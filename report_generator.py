from io import BytesIO
from docx import Document

def generate_report(data):
    """
    data: словарь с ключами:
        - region
        - partner
        - year
        - step_4_value (опционально)
        - step_5_value (опционально)
        - step_6_value (опционально)
    """
    doc = Document()
    doc.add_heading('Отчёт', 0)
    
    doc.add_paragraph(f"Регион: {data.get('region', 'Не указано')}")
    doc.add_paragraph(f"Страна-партнёр: {data.get('partner', 'Не указано')}")
    doc.add_paragraph(f"Год: {data.get('year', 'Не указано')}")
    
    if 'step_4_value' in data:
        doc.add_paragraph(f"Доп. настройка №1: {data['step_4_value']}")
    if 'step_5_value' in data:
        doc.add_paragraph(f"Доп. настройка №2: {data['step_5_value']}")
    if 'step_6_value' in data:
        doc.add_paragraph(f"Доп. настройка №3: {data['step_6_value']}")
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
